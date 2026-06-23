"""Minimal Markdown → .docx converter for manuscript drafts (python-docx).

Handles headings, paragraphs (with **bold**), bullet lists, fenced code blocks,
and GitHub-style tables. Top-of-file guide blockquotes (lines starting with '>')
are skipped so the output is submission-clean. Not a full Markdown parser — meant
for the COMID manuscript drafts only.

Usage:
    python scripts/md_to_docx.py <input.md> <output.docx>
"""

import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# Korean-capable font (Windows default). Set ascii/hAnsi/eastAsia so both Latin
# and Hangul glyphs render; without w:eastAsia, Hangul shows as boxes.
KFONT = "Malgun Gothic"


def _kfont(run):
    """Force a run to use a Korean-capable font for Latin + East-Asian glyphs."""
    run.font.name = KFONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), KFONT)
    rfonts.set(qn("w:hAnsi"), KFONT)
    rfonts.set(qn("w:eastAsia"), KFONT)


def _apply_fonts(doc):
    """Set the Korean-capable font on every paragraph/heading style."""
    for style in doc.styles:
        try:
            style.font.name = KFONT
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), KFONT)
            rfonts.set(qn("w:hAnsi"), KFONT)
            rfonts.set(qn("w:eastAsia"), KFONT)
        except Exception:
            pass


def _add_inline(paragraph, text):
    """Add text with **bold** segments to a paragraph."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            _kfont(run)
        elif part:
            run = paragraph.add_run(part)
            _kfont(run)


def _is_separator(cells):
    return all(set(c) <= set("-: ") and c for c in cells)


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    base_dir = os.path.dirname(os.path.abspath(md_path))
    doc = Document()
    _apply_fonts(doc)
    i, n = 0, len(lines)
    in_code = False
    code_buf = []

    while i < n:
        line = lines[i].rstrip()

        # fenced code block
        if line.strip().startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run._element.get_or_add_rPr().get_or_add_rFonts().set(
                    qn("w:eastAsia"), KFONT)
                code_buf, in_code = [], False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # skip guide blockquotes
        if line.startswith(">"):
            i += 1
            continue

        # image ![alt](path)
        m_img = re.match(r"^!\[(.*?)\]\((.+?)\)\s*$", line.strip())
        if m_img:
            alt, path = m_img.group(1), m_img.group(2)
            full = path if os.path.isabs(path) else os.path.join(base_dir, path)
            if os.path.exists(full):
                doc.add_picture(full, width=Inches(6.0))
                if alt:
                    run = doc.add_paragraph().add_run(alt)
                    run.italic = True
                    run.font.size = Pt(9)
                    _kfont(run)
            else:
                doc.add_paragraph(f"[missing image: {path}]")
            i += 1
            continue

        # table block
        if line.strip().startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = []
            for raw in block:
                cells = [c.strip() for c in raw.strip("|").split("|")]
                if not _is_separator(cells):
                    rows.append(cells)
            if rows:
                ncol = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=ncol)
                table.style = "Table Grid"
                for r in rows:
                    tcells = table.add_row().cells
                    for j in range(ncol):
                        _add_inline(tcells[j].paragraphs[0], r[j] if j < len(r) else "")
            continue

        # headings
        if line.startswith("#### "):
            for r in doc.add_heading(line[5:], level=3).runs:
                _kfont(r)
        elif line.startswith("### "):
            for r in doc.add_heading(line[4:], level=2).runs:
                _kfont(r)
        elif line.startswith("## "):
            for r in doc.add_heading(line[3:], level=1).runs:
                _kfont(r)
        elif line.startswith("# "):
            for r in doc.add_heading(line[2:], level=0).runs:
                _kfont(r)
        elif line.strip() == "---":
            pass
        elif line.startswith("- "):
            _add_inline(doc.add_paragraph(style="List Bullet"), line[2:])
        elif line.strip() == "":
            pass
        else:
            _add_inline(doc.add_paragraph(), line)
        i += 1

    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
